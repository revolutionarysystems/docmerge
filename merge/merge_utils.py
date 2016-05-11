import os
import re
import string
import zipfile
import tempfile
import shutil
from docx import Document
#  from docx.text.paragraph import Paragraph
from django.template import Template, Context
from markdown import markdown
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

from .resource_utils import (
    get_working_dir, strip_xml_dec, get_output_dir, get_local_dir, get_local_txt_content)
from .config import install_name

def replaceParams(txt, subs):
    for key in subs.keys():
        old = "${"+key+"}"
        if (txt.find(old)>=0):
            txt = txt.replace(old, subs[key])
    return txt

def removePara(para):
        p = para._element
        p.getparent().remove(p)
        p._p = p._element = None

def substituteVariablesPlain(fileNameIn, fileNameOut, subs):
    c = Context(subs)
    fileIn = open(fileNameIn, "r")
    fullText = fileIn.read()
    t = Template(fullText)
    xtxt = t.render(c)
    fileOut = open(fileNameOut, "w")
    fileOut.write(xtxt)
    return {"file":fileNameOut}
    
def substituteVariablesPlainString(stringIn, subs):
    c = Context(subs)
    fullText = stringIn
    t = Template(fullText)
    xtxt = t.render(c)
    return xtxt
    
def preprocess(text):
    text = text.replace("{% #A %}", "{% cycle 'A' 'B' 'C' 'D' 'E' 'F' 'G' 'H' 'I' 'J' 'K' 'L' 'M' 'N' 'O' 'P' 'Q' 'R' 'S' 'T' 'U' 'V' 'W' 'X' 'Y' 'Z'%}")    
    text = text.replace("{% #9 %}", "{% cycle '1' '2' '3' '4' '5' '6' '7' '8' '9' '10' '11' '12' '13' '14' '15' '16' '17' '18' '19' '20' '21' '22' '23' '24' '25' '26' '27' '28' '29' '30' as level1 %}")    
    text = text.replace("{% #9= %}", "{{ level1 }}")    
    return text

def apply_sequence(text):
    alf = string.ascii_uppercase
    target = '[% #A %]'
    sub = text.find(target)
    n= 0
    while sub>=1:
        text = text[:sub]+alf[n]+text[sub+len(target):]
        n+=1
        sub = text.find(target)
    return text

def docx_copy_run_style_from(run1, run2):
    run1.font.color.rgb = run2.font.color.rgb
    run1.font.all_caps = run2.font.all_caps
    run1.font.bold = run2.font.bold
    run1.font.italic = run2.font.italic
    run1.font.size = run2.font.size
    run1.font.underline = run2.font.underline
    #complex_script, cs_bold, cs_italic, double_strike, emboss, hidden, highlight_color,
    #imprint, math, name, no_proof, outline, rtl, shadow, small_caps, snap_to_grid, spec_vanish, 
    #strike, superscript, underline, web_hidden

def docx_copy_para_format_from(para1, para2):
    para1.paragraph_format.alignment = para2.paragraph_format.alignment
    para1.paragraph_format.first_line_indent = para2.paragraph_format.first_line_indent
    para1.paragraph_format.keep_together = para2.paragraph_format.keep_together
    para1.paragraph_format.keep_with_next = para2.paragraph_format.keep_with_next
    para1.paragraph_format.left_indent = para2.paragraph_format.left_indent
    try:
        para1.paragraph_format.line_spacing = para2.paragraph_format.line_spacing
        para1.paragraph_format.line_spacing_rule = para2.paragraph_format.line_spacing_rule
    except ValueError:
        pass
    para1.paragraph_format.page_break_before = para2.paragraph_format.page_break_before
    para1.paragraph_format.right_indent = para2.paragraph_format.right_indent
    para1.paragraph_format.space_after = para2.paragraph_format.space_after
    para1.paragraph_format.space_before = para2.paragraph_format.space_before
    para1.paragraph_format.widow_control = para2.paragraph_format.widow_control

def isControlLine(s):
    s = s.split("+")[0]
    s = s.strip()
    if s[:2]=="{%" and s[-2:]=="%}" and s.find("%}")== s.rfind("%}"):
        if s.find("include")>=0:
            return False
        else:
            return True
    else:
        return False

def wrap_list_xml(childlist, root_tag, child_tag):
    out = "<"+root_tag+">\n"
    for child in childlist:
        out+= "\t<"+child_tag+">"+child+"</"+child_tag+">\n"
    out+= "</"+root_tag+">"
    return out

def docx_text(file_name_in):
    doc_in = Document(docx=file_name_in)
    paras=doc_in.paragraphs
    fullText="" 
    for para in paras:
        paraText=para.text
        fullText+=paraText
    return fullText

def extract_regex_matches_docx(file_name_in, regex, wrap=None, root_tag="list", child_tag="item"):
    text = docx_text(file_name_in)
    p = re.compile(regex)
    m=p.findall(text)
    if wrap:
        return wrap_list_xml(m, root_tag, child_tag)
    else:
        return m

    #            literal = m.group('literal')




def substituteVariablesDocx(file_name_in, fileNameOut, subs):
    c = Context(subs)
    doc_in = Document(docx=file_name_in)
    doc_temp = Document()
    paras=doc_in.paragraphs
    fullText="" 
    i = 0
    styles = {}
    for para in paras:
        paraText=""
        p = doc_temp.add_paragraph(style = para.style)
        docx_copy_para_format_from(p, para)
        j = 0
        runs = para.runs
        for run in runs:
            txt = run.text
            paraText+= txt+"+"+str(j)+"+run+"
            r = p.add_run(text = txt, style=run.style)
            docx_copy_run_style_from(r, run)
            j+=1
        fullText+= paraText+str(i)+"+para+"
        i+=1
    fullText = preprocess(fullText)
    t = Template(fullText)
    xtxt = t.render(c)
    xtxt = apply_sequence(xtxt)
    xParaTxts = xtxt.split("+para+")
    for p in paras:
        removePara(p)

    doc_in.paragraphs.clear()
    paras=doc_temp.paragraphs
    for xParaTxt in xParaTxts:
        runTxts = xParaTxt.split("+run+")
        if runTxts[-1]!='':
            para_n = int(runTxts[-1])
            p = doc_in.add_paragraph(style=paras[para_n].style)
            docx_copy_para_format_from(p, paras[para_n])
            for runTxt in runTxts[:-1]:
                try:
                    txt = runTxt.split("+")[-2]
                except:
                    txt=""
                run_n = int(runTxt.split("+")[-1])
                r = p.add_run(text=txt, style=paras[para_n].runs[run_n].style)
                docx_copy_run_style_from(r, paras[para_n].runs[run_n])
            if isControlLine(paras[para_n].text):
                p.text="{}"

    for p in doc_in.paragraphs:
        if p.text=="{}":
            removePara(p)
    doc_in.save(fileNameOut)
    return {"file":fileNameOut}


def print_doc(doc):
    paras=doc.paragraphs
    print("...")
    for para in paras[14:20]:
        print(para.text)
    print("...")

def combine_docx(file_names, file_name_out):
    combined_document = Document(file_names[0])
    count, number_of_files = 0, len(file_names)
    for file in file_names[1:]:
        if file == "pagebreak":
            combined_document.add_page_break()
        else:    
            sub_doc = Document(file)
            for para in sub_doc.paragraphs:
                pnew = combined_document.add_paragraph(style=para.style)
                docx_copy_para_format_from(pnew, para)
                runs = para.runs
                for run in runs:
                    rnew = pnew.add_run(text=run.text, style=run.style)
                    docx_copy_run_style_from(rnew, run)

    combined_document.save(file_name_out)
    return {"file":file_name_out}

def convert_markdown(fileNameIn, fileNameOut):
    fileIn = open(fileNameIn, "r")
    fileOut = open(fileNameOut, "w")
    fileOut.write(markdown(fileIn.read()))
    fileOut.close()
    return {"file":fileNameOut}

def email_file(baseFileName, me, you, subject, credentials):

    msg = MIMEMultipart('alternative')
    msg['Subject'] = subject
    msg['From'] = me
    msg['To'] = you.replace(" ","+")

    mimetypes = ["text/plain", "text/html"]
    file_exts = [".md", ".html"]

    for mime in zip(mimetypes, file_exts):
        file=baseFileName+mime[1]
        fp = open(file, 'r')
        content = fp.read()
        fp.close()
#        print(mime[0].split("/"))
        msg.attach(MIMEText(content, mime[0].split("/")[1]))    
    
    username = credentials["username"]
    password = credentials["password"]
    server = smtplib.SMTP(credentials["server"])

    server.ehlo()
    server.starttls()
    server.ehlo()
    server.login(username,password)
#    print(me)
#    print(you.replace(" ","+"))
#    print(msg.as_string())
    response = server.sendmail(me, [you.replace(" ","+")], msg.as_string())
    server.quit()
    return {"email":you.replace(" ","+")}

def merge_docx_footer(full_local_filename, subs):
    docx_filename = full_local_filename
    f = open(docx_filename, 'rb')
    zip = zipfile.ZipFile(f)
    xml_content = zip.read('word/footer1.xml')
    xml_content = xml_content.decode("ISO-8859-1")
    try:
        xml_content = substituteVariablesPlainString(xml_content, subs)
    except:
        pass
    tmp_dir = tempfile.mkdtemp()
    zip.extractall(tmp_dir)

    with open(os.path.join(tmp_dir,'word/footer1.xml'), 'w') as f:
        f.write(xml_content)
    filenames = zip.namelist()
    zip_copy_filename = docx_filename
    with zipfile.ZipFile(zip_copy_filename, "w") as docx:
        for filename in filenames:
            docx.write(os.path.join(tmp_dir,filename), filename)
    shutil.rmtree(tmp_dir)
    return({"file":docx_filename})
